import os
import hashlib
import tiktoken
from typing import List, Dict, Any, Optional
from pathlib import Path
import PyPDF2
import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
import logging

from config import config

class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.CHUNK_SIZE,
            chunk_overlap=config.CHUNK_OVERLAP,
            length_function=self._count_tokens,  # Use token-based splitting
            separators=["\n\n", "\n", ". ", " ", ""]
        )

        # Initialize token encoder for accurate counting
        try:
            self.encoding = tiktoken.encoding_for_model(config.LLM_MODEL)
        except KeyError:
            self.encoding = tiktoken.get_encoding("cl100k_base")  # Default encoding

        self.supported_extensions = {'.pdf', '.docx', '.txt', '.md'}

        # Set up logging
        logging.basicConfig(level=logging.INFO)
        self.logger = logging.getLogger(__name__)

    def _count_tokens(self, text: str) -> int:
        """Count tokens in text using the model's tokenizer"""
        return len(self.encoding.encode(text))

    def _generate_document_id(self, file_path: str, content: str) -> str:
        """Generate unique ID for document based on path and content"""
        combined = f"{file_path}:{content[:100]}"
        return hashlib.md5(combined.encode()).hexdigest()

    def _extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
        except Exception as e:
            self.logger.error(f"Error reading PDF {file_path}: {e}")
            return ""

    def _extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from Word document"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            self.logger.error(f"Error reading DOCX {file_path}: {e}")
            return ""

    def _extract_text_from_txt(self, file_path: Path) -> str:
        """Extract text from text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read().strip()
            except Exception as e:
                self.logger.error(f"Error reading TXT {file_path}: {e}")
                return ""
        except Exception as e:
            self.logger.error(f"Error reading TXT {file_path}: {e}")
            return ""

    def _clean_text(self, text: str) -> str:
        """Clean and normalize text content"""
        # Remove excessive whitespace
        text = ' '.join(text.split())

        # Remove very short lines that might be noise
        lines = text.split('\n')
        cleaned_lines = [line.strip() for line in lines if len(line.strip()) > 10]

        return '\n'.join(cleaned_lines)

    def extract_text_from_file(self, file_path: Path) -> Optional[str]:
        """Extract text from a single file based on its extension"""
        if not file_path.exists():
            self.logger.error(f"File not found: {file_path}")
            return None

        extension = file_path.suffix.lower()

        if extension == '.pdf':
            text = self._extract_text_from_pdf(file_path)
        elif extension == '.docx':
            text = self._extract_text_from_docx(file_path)
        elif extension in ['.txt', '.md']:
            text = self._extract_text_from_txt(file_path)
        else:
            self.logger.warning(f"Unsupported file type: {extension}")
            return None

        if not text:
            self.logger.warning(f"No text extracted from {file_path}")
            return None

        # Clean the extracted text
        cleaned_text = self._clean_text(text)

        # Check token count
        token_count = self._count_tokens(cleaned_text)
        if token_count > config.token_config.MAX_EMBEDDING_TOKENS:
            self.logger.warning(
                f"Document {file_path} has {token_count} tokens, "
                f"exceeds limit of {config.token_config.MAX_EMBEDDING_TOKENS}"
            )

        return cleaned_text

    def create_chunks(self, text: str, metadata: Dict[str, Any] = None) -> List[Document]:
        """Split text into chunks and create Document objects"""
        if not text:
            return []

        metadata = metadata or {}

        # Create chunks using token-aware splitter
        chunks = self.text_splitter.split_text(text)

        documents = []
        for i, chunk in enumerate(chunks):
            # Skip very small chunks
            if self._count_tokens(chunk) < 50:
                continue

            # Create document with metadata
            doc_metadata = {
                **metadata,
                'chunk_index': i,
                'total_chunks': len(chunks),
                'token_count': self._count_tokens(chunk),
                'chunk_id': hashlib.md5(chunk.encode()).hexdigest()
            }

            doc = Document(
                page_content=chunk,
                metadata=doc_metadata
            )
            documents.append(doc)

        return documents

    def process_single_document(self, file_path: Path) -> List[Document]:
        """Process a single document file"""
        self.logger.info(f"Processing document: {file_path}")

        # Extract text
        text = self.extract_text_from_file(file_path)
        if not text:
            return []

        # Create metadata
        metadata = {
            'source': str(file_path),
            'filename': file_path.name,
            'file_extension': file_path.suffix,
            'file_size': file_path.stat().st_size,
            'document_id': self._generate_document_id(str(file_path), text)
        }

        # Create chunks
        documents = self.create_chunks(text, metadata)

        self.logger.info(
            f"Created {len(documents)} chunks from {file_path} "
            f"(total tokens: {sum(doc.metadata['token_count'] for doc in documents)})"
        )

        return documents

    def process_documents_directory(self, directory_path: str = None) -> List[Document]:
        """Process all documents in the documents directory"""
        directory_path = directory_path or config.DOCUMENTS_PATH
        doc_dir = Path(directory_path)

        if not doc_dir.exists():
            self.logger.error(f"Documents directory not found: {doc_dir}")
            return []

        all_documents = []
        total_tokens = 0

        # Process each supported file
        for file_path in doc_dir.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_extensions:
                documents = self.process_single_document(file_path)
                all_documents.extend(documents)

                # Track token usage
                file_tokens = sum(doc.metadata['token_count'] for doc in documents)
                total_tokens += file_tokens

        self.logger.info(
            f"Processed {len(all_documents)} total chunks "
            f"from {len([f for f in doc_dir.rglob('*') if f.suffix.lower() in self.supported_extensions])} files "
            f"(total tokens: {total_tokens})"
        )

        # Log token usage
        config.save_token_usage(total_tokens)

        return all_documents

    def get_document_stats(self, documents: List[Document]) -> Dict[str, Any]:
        """Get statistics about processed documents"""
        if not documents:
            return {}

        total_tokens = sum(doc.metadata['token_count'] for doc in documents)
        avg_tokens = total_tokens / len(documents)

        file_types = {}
        for doc in documents:
            ext = doc.metadata['file_extension']
            file_types[ext] = file_types.get(ext, 0) + 1

        return {
            'total_chunks': len(documents),
            'total_tokens': total_tokens,
            'average_tokens_per_chunk': round(avg_tokens, 2),
            'file_types': file_types,
            'estimated_embedding_cost': config.calculate_estimated_cost(total_tokens)
        }